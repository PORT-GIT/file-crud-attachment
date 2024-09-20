# utils.py
import docx

def extract_text_from_word(file_path):
    try:
        #opens the word document
        doc = docx.Document(file_path)
        full_text = []

        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            full_text.append(paragraph.text)

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text)  # Get text from each cell in the row
                full_text.append(" | ".join(row_data))  # Combine cell text with a separator

        return "\n".join(full_text)  # Combine all extracted text into one string
    except Exception as e:
        print(f"Error extracting text: {e}")
        return None
